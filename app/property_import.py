"""Bulk property-import parsing: Excel column-mapping/normalization, Word text
extraction (the LLM structuring pass itself lives in voice-service — see
voice-service/app/property_extraction.py), and likely-duplicate detection
against a broker's existing listings. Deliberately free of FastAPI/DB-session
concerns so it's easy to reason about and reuse from both /parse and /commit
in app.routers.property_import.
"""
import re
from dataclasses import dataclass
from io import BytesIO

import docx
import openpyxl

from app.models import PriceUnit, Property, PropertyStatus, PropertyType, SizeUnit
from app.schemas import ImportDuplicateMatch, ImportRowData

# Plenty for a broker's listing sheet without risking a multi-minute parse/LLM pass.
MAX_IMPORT_ROWS = 500
# Bounds the prompt sent to voice-service for Word-doc extraction.
MAX_DOCX_TEXT_CHARS = 40_000

FIELD_ORDER = [
    "type",
    "title",
    "area",
    "city",
    "price",
    "price_unit",
    "size_value",
    "size_unit",
    "bedrooms",
    "amenities",
    "status",
    "description",
]

# Deliberately exact-match (post-normalization), not fuzzy — good column names
# auto-map, anything else falls to the broker via the column-mapping UI rather
# than risking a wrong silent guess.
FIELD_SYNONYMS: dict[str, set[str]] = {
    "type": {"type", "property type", "category"},
    "title": {"title", "name", "listing title", "listing name"},
    "area": {"area", "locality", "neighborhood", "neighbourhood", "location area", "location"},
    "city": {"city", "town"},
    "price": {"price", "rate", "cost", "amount", "asking price"},
    "price_unit": {"price unit", "rate type", "pricing"},
    "size_value": {
        "size",
        "size value",
        "sqft",
        "area sqft",
        "built up area",
        "carpet area",
        "plot size",
        "size sqft",
    },
    "size_unit": {"size unit", "unit"},
    "bedrooms": {"bedrooms", "bhk", "beds", "no of bedrooms"},
    "amenities": {"amenities", "features", "facilities"},
    "status": {"status", "availability"},
    "description": {"description", "details", "notes", "remarks"},
}


class ExcelParseError(ValueError):
    pass


class DocxParseError(ValueError):
    pass


def _normalize_header(header: str) -> str:
    return re.sub(r"[\s_\-]+", " ", header.strip().lower()).strip()


def guess_column_mapping(headers: list[str]) -> dict[str, str]:
    """Best-effort field -> source-header guess. Only fills in fields whose
    normalized header exactly matches a known synonym; everything else is left
    for the broker to map explicitly in the column-mapping UI."""
    normalized = {header: _normalize_header(header) for header in headers}
    mapping: dict[str, str] = {}
    for field, synonyms in FIELD_SYNONYMS.items():
        for header, normalized_header in normalized.items():
            if normalized_header in synonyms:
                mapping[field] = header
                break
    return mapping


def parse_excel(raw_bytes: bytes) -> tuple[list[str], list[dict[str, object]]]:
    """Reads the first worksheet, first row as headers. Returns raw headers plus
    one dict (header -> raw cell value) per non-blank data row, capped at
    MAX_IMPORT_ROWS."""
    try:
        workbook = openpyxl.load_workbook(BytesIO(raw_bytes), read_only=True, data_only=True)
    except Exception as exc:
        raise ExcelParseError("Could not read this file as an Excel (.xlsx) spreadsheet") from exc

    try:
        sheet = workbook.worksheets[0]
        rows_iter = sheet.iter_rows(values_only=True)
        try:
            header_row = next(rows_iter)
        except StopIteration as exc:
            raise ExcelParseError("Spreadsheet is empty") from exc

        headers = [
            str(cell).strip() if cell is not None and str(cell).strip() else f"Column {i + 1}"
            for i, cell in enumerate(header_row)
        ]

        raw_rows: list[dict[str, object]] = []
        for row in rows_iter:
            if row is None or all(cell is None or str(cell).strip() == "" for cell in row):
                continue
            raw_rows.append({headers[i]: row[i] for i in range(min(len(headers), len(row)))})
            if len(raw_rows) >= MAX_IMPORT_ROWS:
                break

        return headers, raw_rows
    finally:
        workbook.close()


def extract_docx_text(raw_bytes: bytes) -> str:
    """Joins paragraph text and table-cell text (many listings get pasted in as
    a table) into one blob for the LLM extraction pass, capped at
    MAX_DOCX_TEXT_CHARS."""
    try:
        document = docx.Document(BytesIO(raw_bytes))
    except Exception as exc:
        raise DocxParseError("Could not read this file as a Word (.docx) document") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    return "\n".join(lines)[:MAX_DOCX_TEXT_CHARS]


# --- Field-level parsing -----------------------------------------------------

_TYPE_MAP: dict[str, PropertyType] = {
    "flat": PropertyType.flat,
    "flats": PropertyType.flat,
    "apartment": PropertyType.flat,
    "apartments": PropertyType.flat,
    "apt": PropertyType.flat,
    "plot": PropertyType.plot,
    "plots": PropertyType.plot,
    "land": PropertyType.plot,
    "open plot": PropertyType.plot,
}

_PRICE_UNIT_MAP: dict[str, PriceUnit] = {
    "total": PriceUnit.total,
    "lump sum": PriceUnit.total,
    "lumpsum": PriceUnit.total,
    "per sqft": PriceUnit.per_sqft,
    "per_sqft": PriceUnit.per_sqft,
    "per sq ft": PriceUnit.per_sqft,
    "persqft": PriceUnit.per_sqft,
    "sqft rate": PriceUnit.per_sqft,
    "rate per sqft": PriceUnit.per_sqft,
}

_SIZE_UNIT_MAP: dict[str, SizeUnit] = {
    "acre": SizeUnit.acres,
    "acres": SizeUnit.acres,
    "sqft": SizeUnit.sqft,
    "sq ft": SizeUnit.sqft,
    "sq. ft.": SizeUnit.sqft,
    "square feet": SizeUnit.sqft,
    "square ft": SizeUnit.sqft,
}

_STATUS_MAP: dict[str, PropertyStatus] = {
    "available": PropertyStatus.available,
    "sold": PropertyStatus.sold,
    "on hold": PropertyStatus.on_hold,
    "on-hold": PropertyStatus.on_hold,
    "onhold": PropertyStatus.on_hold,
    "hold": PropertyStatus.on_hold,
}

_LAKH = 100_000
_CRORE = 10_000_000


def _clean_text(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def _parse_str(value: object) -> str | None:
    return _clean_text(value)


def _parse_type(value: object) -> PropertyType | None:
    text = _clean_text(value)
    return _TYPE_MAP.get(text.lower()) if text else None


def _parse_price_unit(value: object) -> PriceUnit:
    text = _clean_text(value)
    return _PRICE_UNIT_MAP.get(text.lower(), PriceUnit.total) if text else PriceUnit.total


def _parse_size_unit(value: object) -> SizeUnit:
    text = _clean_text(value)
    return _SIZE_UNIT_MAP.get(text.lower(), SizeUnit.sqft) if text else SizeUnit.sqft


def _parse_status(value: object) -> PropertyStatus:
    text = _clean_text(value)
    return _STATUS_MAP.get(text.lower(), PropertyStatus.available) if text else PropertyStatus.available


def _parse_float(value: object) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip()
    if not text:
        return None

    lowered = text.lower().replace(",", "")
    lowered = re.sub(r"[₹$]|rs\.?|inr", "", lowered).strip()

    multiplier = 1.0
    if "crore" in lowered or re.search(r"\bcr\b", lowered):
        multiplier = _CRORE
        lowered = re.sub(r"crores?|\bcr\b", "", lowered).strip()
    elif "lakh" in lowered or re.search(r"\blac\b", lowered):
        multiplier = _LAKH
        lowered = re.sub(r"lakhs?|\blacs?\b", "", lowered).strip()

    match = re.search(r"-?\d+(\.\d+)?", lowered)
    if not match:
        return None
    return float(match.group()) * multiplier


def _parse_int(value: object) -> int | None:
    parsed = _parse_float(value)
    return int(round(parsed)) if parsed is not None else None


def _parse_amenities(value: object) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    if not text:
        return []
    return [part.strip() for part in re.split(r"[,;]", text) if part.strip()]


_PARSERS = {
    "type": _parse_type,
    "title": _parse_str,
    "area": _parse_str,
    "city": _parse_str,
    "price": _parse_float,
    "price_unit": _parse_price_unit,
    "size_value": _parse_float,
    "size_unit": _parse_size_unit,
    "bedrooms": _parse_int,
    "amenities": _parse_amenities,
    "status": _parse_status,
    "description": _parse_str,
}


def normalize_excel_row(raw: dict[str, object], mapping: dict[str, str]) -> ImportRowData:
    values = {}
    for field in FIELD_ORDER:
        header = mapping.get(field)
        raw_value = raw.get(header) if header else None
        values[field] = _PARSERS[field](raw_value)
    return ImportRowData(**values)


def normalize_extracted_row(extracted: dict) -> ImportRowData:
    """LLM-extracted records are already field-keyed (see
    voice-service/app/property_extraction.py's ExtractedProperty) — still run
    through the same parsers as the Excel path so both sources share one
    normalization behavior (price/unit synonyms, amenity splitting, etc.)."""
    values = {field: _PARSERS[field](extracted.get(field)) for field in FIELD_ORDER if field != "status"}
    values["status"] = PropertyStatus.available  # not knowable from free-form text
    return ImportRowData(**values)


def cell_to_display_str(value: object) -> str:
    if value is None:
        return ""
    return str(value)


# --- Validation ---------------------------------------------------------------


def validate_row(data: ImportRowData) -> list[str]:
    errors: list[str] = []
    if data.type is None:
        errors.append("Property type is required (flat or plot)")
    if not data.title:
        errors.append("Title is required")
    if not data.area:
        errors.append("Area/locality is required")
    if not data.city:
        errors.append("City is required")
    if data.price is None or data.price <= 0:
        errors.append("Price is required and must be greater than 0")
    if data.size_value is None or data.size_value <= 0:
        errors.append("Size is required and must be greater than 0")
    return errors


# --- Duplicate detection -------------------------------------------------------

_SIZE_TOLERANCE = 0.15
_PRICE_TOLERANCE = 0.10


@dataclass
class BatchRowRef:
    row_number: int
    data: ImportRowData


def _is_close(a: float, b: float, tolerance: float) -> bool:
    if b == 0:
        return False
    return abs(a - b) / abs(b) <= tolerance


def _same_location(a_area: str, a_city: str, b_area: str, b_city: str) -> bool:
    return a_area.strip().lower() == b_area.strip().lower() and a_city.strip().lower() == b_city.strip().lower()


def _match_reasons(
    candidate_size: float | None,
    candidate_size_unit: SizeUnit,
    candidate_price: float | None,
    candidate_price_unit: PriceUnit,
    other_size: float | None,
    other_size_unit: SizeUnit,
    other_price: float | None,
    other_price_unit: PriceUnit,
) -> list[str] | None:
    reasons = ["same area & city"]
    matched = False
    if (
        candidate_size is not None
        and other_size is not None
        and candidate_size_unit == other_size_unit
        and _is_close(candidate_size, other_size, _SIZE_TOLERANCE)
    ):
        reasons.append("similar size")
        matched = True
    if (
        candidate_price is not None
        and other_price is not None
        and candidate_price_unit == other_price_unit
        and _is_close(candidate_price, other_price, _PRICE_TOLERANCE)
    ):
        reasons.append("similar price")
        matched = True
    return reasons if matched else None


def find_duplicates(
    candidate: ImportRowData,
    existing_properties: list[Property],
    batch_so_far: list[BatchRowRef],
) -> list[ImportDuplicateMatch]:
    if not candidate.area or not candidate.city:
        return []

    matches: list[ImportDuplicateMatch] = []

    for existing in existing_properties:
        if not _same_location(candidate.area, candidate.city, existing.area, existing.city):
            continue
        reasons = _match_reasons(
            candidate.size_value,
            candidate.size_unit,
            candidate.price,
            candidate.price_unit,
            float(existing.size_value),
            existing.size_unit,
            float(existing.price),
            existing.price_unit,
        )
        if reasons:
            matches.append(
                ImportDuplicateMatch(
                    kind="existing",
                    property_id=existing.id,
                    row_number=None,
                    title=existing.title,
                    reason=", ".join(reasons),
                )
            )

    for ref in batch_so_far:
        if not ref.data.area or not ref.data.city:
            continue
        if not _same_location(candidate.area, candidate.city, ref.data.area, ref.data.city):
            continue
        reasons = _match_reasons(
            candidate.size_value,
            candidate.size_unit,
            candidate.price,
            candidate.price_unit,
            ref.data.size_value,
            ref.data.size_unit,
            ref.data.price,
            ref.data.price_unit,
        )
        if reasons:
            matches.append(
                ImportDuplicateMatch(
                    kind="batch",
                    property_id=None,
                    row_number=ref.row_number,
                    title=ref.data.title or f"Row {ref.row_number}",
                    reason=", ".join(reasons),
                )
            )

    return matches


def apply_import_fields(prop: Property, data: ImportRowData) -> None:
    """Full-replace a Property's fields from a finalized import row — same
    semantics as app.routers.properties._apply_fields for the manual-entry PUT."""
    assert data.type is not None
    assert data.title is not None
    assert data.area is not None
    assert data.city is not None
    assert data.price is not None
    assert data.size_value is not None

    prop.type = data.type
    prop.title = data.title
    prop.area = data.area
    prop.city = data.city
    prop.price = data.price
    prop.price_unit = data.price_unit
    prop.size_value = data.size_value
    prop.size_unit = data.size_unit
    prop.bedrooms = data.bedrooms
    prop.amenities = data.amenities
    prop.status = data.status
    prop.description = data.description
